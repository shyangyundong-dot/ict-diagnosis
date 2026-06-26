"""生成《列收模式（27 号文）上线公告》Word 文档。

与 generate.py 同款样式（微软雅黑 + 蓝色标题 + 灰底提示框）。docx 产物入仓库根、gitignored；
改文字只动本脚本（纯文本 diff）。重跑：
    /path/to/python scripts/usage-doc/generate_announcement.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

_HERE = os.path.dirname(os.path.abspath(__file__))
_REPO_ROOT = os.path.abspath(os.path.join(_HERE, "..", ".."))
DOC_PATH = os.path.join(_REPO_ROOT, "ICT合规诊断工具-列收模式上线公告.docx")

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.54)
section.top_margin = section.bottom_margin = Cm(2.2)


def set_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), "微软雅黑")
        rPr.insert(0, rFonts)
    return p


def body(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_font(run, size=11, bold=bold)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def add_tip(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF7E6')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10, color=(0x92, 0x40, 0x0e))
    doc.add_paragraph()


# ── 标题 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ICT 合规诊断工具更新公告")
set_font(run, size=20, bold=True, color=(0x1a, 0x56, 0xdb))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("列收模式（集团 27 号文）上线")
set_font(run, size=14, color=(0x37, 0x51, 0x8c))

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("广州电信云中台 · 2026 年 6 月")
set_font(run, size=10, color=(0x6b, 0x72, 0x80))

doc.add_paragraph()
body("各位同事：")
body("诊断工具已按集团《关于 ICT 业务收入口径相关工作的通知》（办政企〔2026〕27 号）完成升级，"
     "即日起在试运行环境（http://183.131.86.84:8090）生效。本次是收入列收口径的重要调整，请留意以下变化。")

h2("一、核心变化：硬件不再「一律不列收」")
body("过去「设备/施工铁律不列收」的旧口径已调整为：")
bullet("集团白名单内的标准化设备/成品软件，在「项目门槛（金额/利润率）+ 控制权」齐备时，可全额列收；")
bullet("不在白名单、或门槛/控制权不达标的，仍按净额；")
bullet("施工类默认净额（施工本质不进白名单全额）。")

h2("二、填报时新增两处要填")
body("1）核算单元卡片 →「集团白名单」：对「设备/标品」单元，请如实标注是否属于 27 号文白名单"
     "标准化设备/成品软件（是/否/不确定；不确定时系统从严，确认后或可改善）。", bold=False)
body("2）「信息解析」面板 →「列收模式信息」段（共 6 项）：是否重大整合、项目整体利润率、付款节点、"
     "产权转移、集采比例、是否资本投资。多为售中/财务信息，AI 常抽不全，需手动确认。")
add_tip("注意：单元「是否列收」不再手工勾选，由系统按列收模式自动算出。")

h2("三、报告新增「列收模式判定」板块")
body("提交诊断后，报告会给出：判定的列收模式（服务整合全额 / 单一履约白名单全额 / 净额 / 资本投资）、"
     "关键占比、资格闸逐项 ✅/❌、每个核算单元的全额/净额结论，以及硬否决（红）/ 软提示（黄）。")

h2("四、判定为「净额」但你认为该全额？")
body("板块会列出「硬否决」原因（如控制权未自证、付款节点不符）和需举证/补正的「软提示」。"
     "可补齐控制权角色、修正字段后重新提交；如确有依据维持全额，按板块提示准备举证材料，由审核人员据实判断。"
     "工具坚持「标风险、举证定生死，不替审核人定罪」。")

h2("五、配套")
bullet("《使用说明》已更新（新增列收模式信息、列收模式判定解读、Q9/Q10），规则库 v1.8。")
bullet("当前为试运行，欢迎试用并反馈：判定结果与理解不一致、白名单归类拿不准、字段填报有疑问，随时截图反馈。")

doc.add_paragraph()
body("如有问题请联系 [负责人/联系方式]。")
doc.add_paragraph()
sign = doc.add_paragraph()
sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sign.add_run("——广州电信云中台")
set_font(run, size=11, color=(0x37, 0x51, 0x8c))

doc.save(DOC_PATH)
print(f"公告已生成：{DOC_PATH}")
