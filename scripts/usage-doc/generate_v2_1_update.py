"""生成《ICT 项目合规诊断工具 v2.1.0 升级说明》Word 文档。

文档产物写入仓库根目录并由 .gitignore 排除；本脚本作为可复现源文件入库。

运行：
    python scripts/usage-doc/generate_v2_1_update.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
DOC_PATH = os.path.join(REPO_ROOT, "ICT项目合规诊断工具-v2.1.0升级说明.docx")

# launch_messaging_guide（compact_reference_guide 派生）设计令牌。
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
PAGE_MARGIN = Inches(1)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

# 中文发布文档的命名字体覆盖：该字体同时覆盖中英文，避免 LibreOffice
# 忽略 eastAsia 字体后出现缺字方框；Windows Word 缺失时会按系统字体回退。
ASCII_FONT = "Hiragino Sans GB"
EAST_ASIA_FONT = "Hiragino Sans GB"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "FFF7E6"
BORDER = "CBD5E1"
WHITE = "FFFFFF"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, *, size=None, bold=None, color=None, italic=None):
    run.font.name = ASCII_FONT
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    rfonts.set(qn("w:cs"), ASCII_FONT)


def set_style_font(style, *, size, bold=False, color="000000"):
    style.font.name = ASCII_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = rgb(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    rfonts.set(qn("w:cs"), ASCII_FONT)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN
    section.left_margin = PAGE_MARGIN
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        set_style_font(style, size=size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("ICT 项目合规诊断工具  ·  升级说明")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run("广州电信云中台  ·  第 ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)
    run = fp.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("ICT 项目合规诊断工具")
    set_run_font(run, size=25, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("v2.1.0 升级说明")
    set_run_font(run, size=15, color=BLUE)

    metadata = [
        ("发布单位", "广州电信云中台"),
        ("发布日期", "2026 年 7 月 20 日"),
        ("版本口径", "规则版本 v2.1.0 · 核算结构 schema v2 · 材料目录 2026-07-19"),
        ("适用范围", "升级后新建并提交的诊断；历史报告保持原规则版本和原始快照"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(f"{label}：")
        set_run_font(run, size=10.5, bold=True, color=INK)
        run = p.add_run(value)
        set_run_font(run, size=10.5, color="334155")

    add_callout(
        doc,
        "本次升级把列收判断从项目级单值进一步细化到最终核算单元：先确认合同结构和拟列收方式，再逐单元完成六到位、R08 控制权、27 号文配套条件及硬转服务核查。系统负责标示风险和列出证据缺口，不替代审核人员作最终认定。",
        fill=CALLOUT,
        accent=BLUE,
    )


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text[: len(bold_prefix)], text[len(bold_prefix) :]
        run = p.add_run(first)
        set_run_font(run, bold=True, color=INK)
        run = p.add_run(rest)
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def create_numbering(doc: Document, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "●" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    ppr.append(indent)
    level.append(ppr)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), ASCII_FONT)
    fonts.set(qn("w:hAnsi"), ASCII_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    rpr.append(fonts)
    level.append(rpr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, *, num_id: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)
    run = p.add_run(text)
    set_run_font(run)
    return p


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(
    table,
    widths: list[int],
    *,
    indent=TABLE_INDENT_DXA,
    repeat_header=False,
):
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if repeat_header and row_index == 0:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_fill(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, *, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def format_cell_text(cell, *, bold=False, color="334155", size=10):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def add_overview_table(doc: Document):
    rows = [
        ("核算结构", "支持同一 BPM 内部分组合、部分独立；PO1–PO4 逐候选组合确认。", "先确认合同结构，再进入列收自检。"),
        ("列收判断", "分开记录拟列收意图、规则派生结果和确认/暂定状态。", "不再用一个项目级结论覆盖全部单元。"),
        ("六到位与 R08", "项目共性事实复用，结论落到每个拟全额最终核算单元。", "某单元失败只影响该单元，不连带其他单元。"),
        ("硬转服务", "对拟全额服务明细逐项检查硬件/施工实质，命中后列证据。", "只标嫌疑，不自动改类型、拆组合或改净额。"),
        ("材料与报告", "材料统一引用主目录，按四类去重；报告展示意图、结果、状态和缺口。", "更便于按材料编号准备线下审核佐证。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("升级模块", "主要变化", "对使用人的影响")
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_border(cell)
        format_cell_text(cell, bold=True, color=INK, size=10)
    for item in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(item):
            cells[idx].text = value
            set_cell_fill(cells[idx], WHITE)
            set_cell_border(cells[idx])
            format_cell_text(cells[idx])
    set_table_geometry(table, [1650, 4290, 3420], repeat_header=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("表 1｜v2.1.0 用户可见变化概览")
    set_run_font(run, size=9, color=MUTED, italic=True)


def add_callout(doc: Document, text: str, *, fill: str, accent: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_fill(cell, fill)
    set_cell_border(cell, color=accent, size="8")
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    format_cell_text(cell, color=INK, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def build_document():
    doc = Document()
    configure_document(doc)
    bullet_num = create_numbering(doc, kind="bullet")
    step_num = create_numbering(doc, kind="decimal")
    add_title_block(doc)

    add_heading(doc, "一、升级概览", 1)
    add_body(
        doc,
        "v2.1.0 延续“工具标风险、举证定生死，不替审核人定罪”的原则，重点解决旧流程中项目级结论过粗、组合关系无法局部表达、缺证据容易被直接当失败，以及材料清单重复维护等问题。",
    )
    add_overview_table(doc)

    add_heading(doc, "二、填报流程有哪些变化", 1)
    add_heading(doc, "2.1 新诊断按六步完成确认", 2)
    for text in (
        "确认原始核算单元：核对 AI 切分的业务块、类型、金额及明细事实。",
        "确认候选组合：对每个候选组合填写 PO1–PO4，参考系统建议后确认组合或分别核算。",
        "确认列收意图：每个非标品最终核算单元选择拟全额或拟净额；标品恒全额。",
        "确认项目共性事实：填写关键角色、交付模式等可被多个单元复用的客观事实。",
        "完成单元自查：仅对拟全额单元逐项确认六到位和 R08；拟净额单元跳过。",
        "补充项目口径政策信息：填写整体利润率、客户类型、付款节点等，系统重新计算 60%/80% 等门槛。",
    ):
        add_list_item(doc, text, num_id=step_num)

    add_heading(doc, "2.2 核算单元类型和白名单", 2)
    for text in (
        "原始单元类型统一为设备、成品软件、施工、服务、标品、其他；“其他”仅是草稿占位，提交前必须重新归类。",
        "设备和成品软件使用白名单三态：是、否、不确定；不确定不会直接按失败处理，但拟全额结果会标为暂定、高风险。",
        "标品指电话、宽带、天翼云等电信自有标准化产品，恒为确认全额，不参与 27 号文白名单判断。",
        "服务侧毛利率只描述应列收/服务侧毛利；项目整体利润率包含硬件等整体口径，两者不可混填。",
    ):
        add_list_item(doc, text, num_id=bullet_num)

    add_heading(doc, "2.3 列收结论改为“三件套”", 2)
    add_body(doc, "每个最终核算单元的报告同时展示三类信息：")
    for text in (
        "列收意图：用户确认的拟全额或拟净额。",
        "列收结果：规则派生的全额或净额，用户和 AI 都不能直接改写。",
        "结果状态：confirmed 表示已确认；provisional 表示暂按当前金额口径测算，但证据或资格尚未证实。",
    ):
        add_list_item(doc, text, num_id=bullet_num)
    add_callout(
        doc,
        "示例：拟全额的设备单元若白名单状态为“不确定”，系统保留“暂按全额测算”，同时标记高风险并列出补证材料；若白名单明确为“否”，则该单元确认净额。两种情况都不会改写其他核算单元的结果。",
        fill=CAUTION,
        accent="D97706",
    )

    add_heading(doc, "三、单元级自查有哪些变化", 1)
    add_heading(doc, "3.1 六到位合并为一套自查", 2)
    add_body(
        doc,
        "原项目级角色检查与服务场景证据现统一归入六到位。关键角色、交付模式等事实只采集一次，但六个维度和综合等级分别落到每个拟全额最终核算单元。系统可生成建议，最终仍由用户确认。",
    )
    for text in (
        "“强”要求六维全部到位或合法不适用、关键角色齐备，并且不存在全部外部交付等冲突信号。",
        "“中”或“无”、适用维度明确不到位、关键角色明确不齐备，会使对应单元确认净额。",
        "“待补证据”不等于失败：该单元仍可暂按全额测算，但必须标为暂定和高风险。",
        "“不适用”只向采购自主、运维自主开放，并且需要无外部采购或无运维/售后义务的事实基础。",
    ):
        add_list_item(doc, text, num_id=bullet_num)

    add_heading(doc, "3.2 R08 控制权独立判断", 2)
    add_body(
        doc,
        "六到位先判断目标核算单元是否具备争取全额的主控和验收主责基础；通过后，R08 再判断主要责任人/代理人会计控制权。二者可复用共性证据，但不合并结论。",
    )

    add_heading(doc, "3.3 硬转服务采用举证式提示", 2)
    add_body(
        doc,
        "当服务明细拟全额且呈现零毛利平进平出、物流供应商直发、无自有能力等硬件/施工实质时，系统标记“硬转服务嫌疑”并生成标准证据清单。组合后的服务明细仍逐项检查。",
    )
    add_callout(
        doc,
        "重要边界：命中嫌疑不等于已经定性。系统不会自动修改业务类型、拆分组合关系或改成净额；应由填报人核对实质并准备材料，由审核人员据证认定。",
        fill=CALLOUT,
        accent=BLUE,
    )

    add_heading(doc, "四、报告和材料清单有哪些变化", 1)
    for text in (
        "“列收模式判定”按最终核算单元展示列收意图、派生结果、确认/暂定状态、原因和证据缺口。",
        "“六到位自查”和“R08 控制权”只展示拟全额单元；没有拟全额单元时相应模块隐藏，不额外抬高风险。",
        "“硬转服务嫌疑”展示命中的原始服务明细和需核验的事实，不替代正式认定。",
        "材料清单统一从材料主目录生成，规则卡和嫌疑卡只显示材料编号，完整组成只在统一清单展开。",
        "材料按基础过程材料、条件性合规材料、财务列收材料、异常补正材料四类归集，并按材料编号去重。",
        "整体风险取规则和单元风险的最高等级，但风险汇总不会反向改写任何其他单元的列收结果。",
    ):
        add_list_item(doc, text, num_id=bullet_num)

    add_heading(doc, "五、变更后的重新确认规则", 1)
    for text in (
        "原始单元类型、组成或组合关系变化：只使受影响最终核算单元的列收意图、六到位和 R08 确认失效。",
        "金额变化：保留单元级人工确认，但会按整个 BPM 重新计算占比、整体利润率等项目口径门槛。",
        "关键角色、交付模式等项目共性事实变化：所有拟全额单元的六到位和 R08 进入需重新确认状态。",
        "新会话保存 schema v2 结构；旧数组和旧 listed 字段仅用于历史报告兼容，不追溯改写。",
    ):
        add_list_item(doc, text, num_id=bullet_num)

    add_heading(doc, "六、使用提醒与工具边界", 1)
    for text in (
        "AI 切分、字段抽取和六到位建议都是草稿；最终核算结构、列收意图及人工自查项必须由用户核对。",
        "暂定全额允许生成自检报告，但不代表正式送审或审核通过；正式审核前应补齐报告列明的证据。",
        "系统不提供附件上传、补证完成跟踪、送省按钮、送审状态或审批流，相关核验在线下完成。",
        "诊断结论仅作风险提示，不替代 BPM 审批、财务判断和审核人员的最终认定。",
    ):
        add_list_item(doc, text, num_id=bullet_num)

    doc.add_page_break()
    add_heading(doc, "七、版本信息", 1)
    version_rows = [
        ("系统升级标识", "v2.1.0"),
        ("规则版本", "v2.1.0（35 条，R01–R37，跳号 R04/R33）"),
        ("核算结构", "schema_version = 2"),
        ("材料目录版本", "2026-07-19"),
        ("材料依据", "《附件2 ICT项目全流程控制权角色职责和佐证材料》（2026-07-18 更新版）"),
        ("发布日期", "2026-07-20"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in version_rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_fill(cells[0], LIGHT_GRAY)
        set_cell_fill(cells[1], WHITE)
        set_cell_border(cells[0])
        set_cell_border(cells[1])
        format_cell_text(cells[0], bold=True, color=INK)
        format_cell_text(cells[1])
    set_table_geometry(table, [2700, 6660])

    add_callout(
        doc,
        "本说明以 2026-07-20 工作区实现及规则库 v2.1.0 为准。条款原文和集团白名单可能后续更新，具体认定仍以集团、省公司最新文件及审核意见为准。",
        fill=CALLOUT,
        accent=BLUE,
    )

    doc.core_properties.title = "ICT 项目合规诊断工具 v2.1.0 升级说明"
    doc.core_properties.subject = "核算单元 v2、单元级列收自检与材料目录升级"
    doc.core_properties.author = "广州电信云中台"
    doc.core_properties.keywords = "ICT, 合规诊断, v2.1.0, 核算单元, 列收"
    doc.save(DOC_PATH)
    print(f"升级说明已生成：{DOC_PATH}")


if __name__ == "__main__":
    build_document()
