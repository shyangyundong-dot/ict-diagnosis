from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 路径相对脚本自身，clone 后可直接跑（无需改绝对路径）
_HERE = os.path.dirname(os.path.abspath(__file__))
_REPO_ROOT = os.path.abspath(os.path.join(_HERE, "..", ".."))
OUT = os.path.join(_HERE, "screenshots")
DOC_PATH = os.path.join(_REPO_ROOT, "ICT项目合规诊断工具-使用说明.docx")

doc = Document()

# ── 页面设置 A4 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.54)
section.top_margin  = section.bottom_margin = Cm(2.54)

# ── 样式辅助 ──
def set_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), "微软雅黑")
        rPr.insert(0, rFonts)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), "微软雅黑")
        rPr.insert(0, rFonts)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(12)
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), "微软雅黑")
        rPr.insert(0, rFonts)
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_img(path, width=Inches(5.5), caption=None):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(12)
            for run in cp.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

def add_tip(text):
    """灰色提示框用表格模拟"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置背景色
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EFF6FF')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10, color=(0x1e, 0x40, 0xaf))
    doc.add_paragraph()

# ════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("ICT 项目合规诊断工具")
set_font(run, size=24, bold=True, color=(0x1a, 0x56, 0xdb))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("使用说明")
set_font(run, size=18, color=(0x37, 0x51, 0x8c))

doc.add_paragraph()
org_p = doc.add_paragraph()
org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = org_p.add_run("广州电信云中台 · 内部工具")
set_font(run, size=12, color=(0x6b, 0x72, 0x80))

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("2026 年 6 月")
set_font(run, size=11, color=(0x6b, 0x72, 0x80))

doc.add_page_break()

# ════════════════════════════════════════════
# 目录（手工列出）
# ════════════════════════════════════════════
h1("目  录")
toc_items = [
    ("一、工具简介", "3"),
    ("二、登录与账号", "3"),
    ("三、发起诊断——对话式信息收集", "4"),
    ("四、信息解析与核算单元确认", "5"),
    ("五、诊断报告解读", "6"),
    ("六、诊断列表", "8"),
    ("七、按 BPM 编号查询", "8"),
    ("八、填报溯源", "9"),
    ("九、管理员后台（仅 admin）", "9"),
    ("十、常见问题", "12"),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{item}")
    set_font(run, size=11)

doc.add_page_break()

# ════════════════════════════════════════════
# 一、工具简介
# ════════════════════════════════════════════
h1("一、工具简介")
body("ICT 项目合规诊断工具是广州电信云中台内部的合规风险自查平台。使用人员通过自然语言描述项目情况，系统自动提取结构化信息、匹配合规规则库，生成个性化风险诊断报告，并给出整改建议和审计材料清单。")
doc.add_paragraph()
body("核心能力：")
bullet("AI 对话提取字段：无需逐项填表，像聊天一样描述项目，AI 自动识别并提取关键合规字段")
bullet("核算单元切分：将一笔合同中的设备、施工、服务等不同业务块分别进行合规判断，避免混算误报")
bullet("规则引擎诊断：覆盖广东电信「六到位核查清单」六个维度，共 35 条规则自动检测")
bullet("AI 个性化分析：针对每条触发规则给出具体分析与整改建议，而非通用模板")
bullet("PDF 报告下载：完整报告可导出存档")
bullet("多角色权限管理：员工、线条主管（reviewer）、管理员三级权限隔离")
doc.add_paragraph()
add_tip("提示：工具定位是「标风险、举证定生死」，输出结果是风险提示，最终合规结论由审核人员依据材料判断，工具不替审核人定罪。")

doc.add_page_break()

# ════════════════════════════════════════════
# 二、登录与账号
# ════════════════════════════════════════════
h1("二、登录与账号")

h2("2.1 登录")
body("在浏览器打开系统地址，输入账号和密码后点击「登录」。账号由管理员创建并分配。")
doc.add_paragraph()
add_img(f"{OUT}/01_login.png", width=Inches(3.5), caption="图 1  登录页面")

h2("2.2 账号角色说明")
body("系统分三级角色，不同角色可访问的数据范围不同：")
doc.add_paragraph()

# 角色表格
tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["角色", "数据范围", "特殊权限"]
rows_data = [
    ("管理员（admin）", "全部诊断记录（含历史存量）", "账号/线条管理、审计日志、存量认领"),
    ("主管（reviewer）", "本线条所有员工的记录 + 自己的", "可提交人工复核结论"),
    ("员工（user）", "仅自己创建的记录", "—"),
]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_font(run, size=11, bold=True)
for ri, row_data in enumerate(rows_data):
    for ci, val in enumerate(row_data):
        cell = tbl.rows[ri+1].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(val)
        set_font(run, size=10)
doc.add_paragraph()

h2("2.3 修改密码")
body("点击右上角「修改密码」可自行修改密码。首次登录时如系统要求强制改密，需先完成改密才能进入其他页面。")
body("如忘记密码，联系管理员在「账号管理」页面重置。")

doc.add_page_break()

# ════════════════════════════════════════════
# 三、发起诊断——对话式信息收集
# ════════════════════════════════════════════
h1("三、发起诊断——对话式信息收集")

h2("3.1 新建诊断")
body("点击首页右上角「+ 新建诊断」，清空当前会话，开始新一轮信息收集。")
body("在底部输入框，用自然语言描述项目情况，例如：")
add_tip("「我有一个给番禺某国企做的系统集成项目，预算 500 万，后向供应商还没定，毛利大概 4% 左右，涉及设备采购和软件服务，打算全额列收。」")
body("AI 助手会根据描述提问补充信息，多轮对话逐步完善所有关键字段。每次输入后按 Enter 或点击发送按钮发送。")
doc.add_paragraph()
add_img(f"{OUT}/02_chat.png", width=Inches(5.8), caption="图 2  对话式信息收集界面")

h2("3.2 右侧「信息解析」面板")
body("对话过程中，右侧面板会实时展示 AI 已识别的字段及其当前值。")
bullet("绿色字段：已识别并确认")
bullet("「待补充 N 项」：仍有必填字段未收集到，继续补充描述")
body("如发现字段识别有误，可直接在面板中点击字段修改。")
doc.add_paragraph()
add_tip("提示：提交诊断前，务必确认「项目类型」字段已正确填写——这是必填项，决定哪些规则会被触发。")

doc.add_page_break()

# ════════════════════════════════════════════
# 四、信息解析与核算单元确认
# ════════════════════════════════════════════
h1("四、信息解析与核算单元确认")

h2("4.1 什么是核算单元")
body("同一笔合同（同一个 BPM 商机）内，可能包含设备采购、施工、服务等多个被分别核算的业务块，每块称为一个「核算单元」。不同类型的核算单元适用不同的合规规则：")
bullet("设备 / 施工类：铁律不列收，合规排除，系统不予报警")
bullet("服务类（且列收）：是合规审查的重点，会触发相关规则检测")
doc.add_paragraph()
body("核算单元切分的意义在于：避免把设备的低毛利混入服务侧毛利，导致误报；同时精准识别「申报为服务、实质为硬件」的包装风险（硬转服务）。")

h2("4.2 AI 切分草稿 → 用户确认")
body("信息收集完成后，系统会在对话区提示「建议切分核算单元」。点击后 AI 自动生成切分草稿，展示在右侧面板的「核算单元」区域。")
body("请逐项核对各单元的：申报类型、金额、毛利、是否列收、是否有电信自有能力。如有误差可直接修改，确认无误后点击「确认核算单元」。")
doc.add_paragraph()
add_tip("提示：核算单元确认后会随诊断一起落库，报告中的「已排除列收」「硬转服务嫌疑」等板块都基于此数据生成。")

h2("4.3 提交诊断")
body("信息收集完毕、核算单元确认后，点击右侧底部「提交诊断」按钮。系统将：")
bullet("第一步：规则引擎立即运行，判定哪些规则被触发")
bullet("第二步：AI 对每条触发规则进行个性化分析（并发执行）")
body("整个过程通常需要 30～90 秒，请耐心等待页面跳转到报告。")

h2("4.4 未切分核算单元时的提示")
body("如果项目涉及设备采购、系统集成等通常包含硬件/施工业务块的类型，但你没有切分核算单元就直接提交，报告顶部会出现一条黄色提示横幅：")
doc.add_paragraph()
add_img(f"{OUT}/11_unit_warning_banner.png", width=Inches(5.8), caption="图 · 未切分核算单元的黄色提示横幅")
doc.add_paragraph()
body("这条提示的含义是：由于没有切分核算单元，系统无法对硬件/施工做「铁律不列收」排除，也无法对服务单元做「硬转服务」举证式检测，本次诊断按项目整体的单一口径给出，结论可能偏严（更容易报风险）。")
body("它不会阻止你提交，诊断报告照常生成；但如果你看到这条横幅，建议返回「信息解析」面板把核算单元切分好、确认后重新提交，得到更精准的诊断。")
doc.add_paragraph()
add_tip("提示：纯服务、纯软件这类通常只有单一业务块的项目，不切分核算单元不会出现此提示，属正常。")

h2("4.5 控制权角色自查（总额法资格）")
body("在「核算单元」段下方，「信息解析」面板还有一段「控制权角色」。这是按广东电信省公司《产数ICT业务高质量发展专项部署材料》的官方框架做的自查——电信在本项目占据哪些「关键角色」（决策/主导/责任），决定项目能不能按总额法列收。")
doc.add_paragraph()
add_img(f"{OUT}/13_panel_control_roles_section.png", width=Inches(4.5), caption="图 · 信息解析面板的控制权角色段")
doc.add_paragraph()
body("勾选规则：")
bullet("「必选」三项（应标签约统筹 6 / 软硬件采购决策 7 / 全流程交付管理与质量责任 9）每项都要——电信在这三个环节都必须是主责者")
bullet("如项目涉及硬件，还要勾「到货验收及设备管理 16」（系统会根据你切的核算单元里有没有「设备/施工」类型自动显示这一项）")
bullet("三组「二选一」（方案 / 交付实施方案 / 实施开发）每组至少占一个——表示电信在该环节的决策权")
doc.add_paragraph()
body("必选都占齐 + 三组各占一个 = 总额法资格成立。这与官方 8 种合法情形对应。")
doc.add_paragraph()
add_tip("提示：AI 通常解析不出这些角色（商机/财务对话里没有「谁干什么」），所以多数情况下需要你手动勾选。可以不填——但项目从字段上看明显奔全额列收时，不勾会在报告里收到「控制权未自证」的中风险提示。")

doc.add_page_break()

# ════════════════════════════════════════════
# 五、诊断报告解读
# ════════════════════════════════════════════
h1("五、诊断报告解读")
add_img(f"{OUT}/08_report.png", width=Inches(5.8), caption="图 3  诊断报告示例")
doc.add_paragraph()

h2("5.1 整体风险等级")
body("报告顶部显示本次诊断的综合风险等级：")
bullet("高风险（红色）：触发了严重合规问题，需立即处理")
bullet("中风险（橙色）：存在需关注的合规隐患")
bullet("低风险（绿色）：未发现明显问题，但仍需关注人工核查项")

h2("5.2 已触发规则")
body("列出系统自动触发的规则，每条规则包含：")
bullet("规则名称与风险等级")
bullet("AI 个性化分析说明：结合项目实际情况解读为何触发")
bullet("整改建议：具体的合规修正方向")
bullet("所需审计材料：需准备哪些证明文件")

h2("5.3 已排除列收（核算单元说明）")
body("如项目包含设备或施工核算单元，报告中会显示「已排除列收」板块，说明这些单元已正确归类为不列收，相关规则（R24/R25/R26）已被系统自动抑制，不计入风险等级。")

h2("5.4 硬转服务嫌疑")
body("如某个申报为「服务」的核算单元呈现以下特征，系统将标记为「硬转服务嫌疑」：")
bullet("零毛利平进平出")
bullet("由供应商物流直发，电信未实质参与")
bullet("无电信自有服务能力")
body("系统不自动定性为违规，而是列出需要举证的材料清单，由审核人员综合判断。")

h2("5.5 控制权角色自查（总额法资格）")
body("报告会根据你在面板里勾选的「控制权角色」呈现一个独立板块，4 种颜色对应 4 种状态：")
bullet("🟢 绿色「✅ 总额法资格成立」——必选都占齐 + 三组各占一个，符合官方 8 种合法情形之一")
bullet("🔴 红色「❌ 总额法资格不成立」——缺关键角色，定性倾向代理人/净额；如维持全额列收，须举证缺失角色到位（计入高风险）")
bullet("🟡 黄色「⚠️ 控制权未自证」——你没勾角色但项目从字段上看明显奔全额列收（如能力充足或服务自有/混合交付）；计入中风险，建议回面板补勾")
bullet("⚪ 灰色「ⓘ 未参与判定」——没勾角色 + 项目本不奔全额（如纯设备销售）；只留痕、不计入风险")
doc.add_paragraph()
add_img(f"{OUT}/12_ctrl_card_ineligible.png", width=Inches(5.8), caption="图 · 控制权角色自查（红色「资格不成立」示例，列出缺失角色）")
doc.add_paragraph()
add_tip("提示：这套自查依据省公司《产数ICT业务高质量发展专项部署材料》的官方 19 角色/8 情形矩阵。和 5.4 硬转服务嫌疑互补——前者是项目级（电信是否主导整个项目），后者是单元级（某个服务单元是否真服务）。两者都报不算重复，从不同尺度看控制权。")

h2("5.6 人工核查项目")
body("部分规则（如「十个不准」、「虚假项目」等）需要人工逐条核对，系统无法自动判断，统一列在「人工核查项目」板块，请自行对照核查。")

h2("5.7 审计材料清单")
body("报告末尾汇总了本次诊断所需的全部审计证明材料，方便统一准备存档。")

h2("5.8 下载报告")
body("点击报告右上角「下载报告」按钮，可将完整报告下载为 PDF 文件。")

h2("5.9 人工复核（主管 / 管理员）")
body("主管（reviewer）和管理员可在报告页面点击「人工复核」，填写复核结论，供留档和追溯使用。员工角色无此操作权限。")

doc.add_page_break()

# ════════════════════════════════════════════
# 六、诊断列表
# ════════════════════════════════════════════
h1("六、诊断列表")
body("点击顶部导航「诊断列表」，可查看所有历史诊断记录。")
add_img(f"{OUT}/03_diagnoses.png", width=Inches(5.8), caption="图 4  诊断列表页面")
doc.add_paragraph()
body("列表按创建时间倒序排列，每行显示：BPM 商机编号、项目类型、整体风险等级、诊断时间、提交人。点击任意一行进入对应报告详情。")
body("显示范围根据角色自动过滤：员工只看自己的，主管看本线条全部，管理员看全部。")

doc.add_page_break()

# ════════════════════════════════════════════
# 七、按 BPM 编号查询
# ════════════════════════════════════════════
h1("七、按 BPM 编号查询")
body("点击顶部导航「BPM 查询」，可输入 BPM 商机编号精确查找历史诊断记录。")
add_img(f"{OUT}/04_bpm_lookup.png", width=Inches(5.8), caption="图 5  BPM 查询页面")
doc.add_paragraph()
body("输入时大小写均可（系统内部统一转大写匹配），点击「查询」后列出该 BPM 编号下的所有诊断记录，方便对同一项目的多次诊断进行对比。")

doc.add_page_break()

# ════════════════════════════════════════════
# 八、填报溯源
# ════════════════════════════════════════════
h1("八、填报溯源")
body("点击顶部导航「填报溯源」，输入诊断 ID 可查看该次诊断的完整填报记录：")
add_img(f"{OUT}/09_traceability.png", width=Inches(5.8), caption="图 6  填报溯源页面")
doc.add_paragraph()
bullet("提交时各字段的结构化取值（不含原始自由文本中的项目名称、客户名等敏感信息）")
bullet("原始对话记录快照（供审核追溯）")
body("此功能主要供主管和管理员在复核或审计时使用，用于核实填报依据。")

doc.add_page_break()

# ════════════════════════════════════════════
# 九、管理员后台
# ════════════════════════════════════════════
h1("九、管理员后台（仅 admin）")
body("管理员账号登录后，顶部导航栏会显示「线条管理」「账号管理」「存量认领」「审计日志」四个管理入口，普通员工和主管不可见。")

h2("9.1 账号管理")
add_img(f"{OUT}/05_admin_users.png", width=Inches(5.8), caption="图 7  账号管理页面")
doc.add_paragraph()
body("账号管理页面可执行：")
bullet("新建账号：填写用户名、姓名、角色（主管/员工）、所属线条，系统生成初始密码，用户首次登录后必须修改")
bullet("编辑账号：修改角色、所属线条")
bullet("重置密码：为用户重置密码，重置后用户下次登录需强制改密")
bullet("禁用 / 启用：禁用账号后该用户立即无法登录（即时生效）")
body("注意：删除操作使用「软删除」（禁用），不会物理删除账号，以保留审计记录完整性。")

h2("9.2 线条管理")
add_img(f"{OUT}/06_admin_lines.png", width=Inches(5.8), caption="图 8  线条管理页面")
doc.add_paragraph()
body("线条是数据隔离的基本单位（对应各业务线条或团队）。管理员可：")
bullet("新建线条：填写线条名称")
bullet("编辑线条：修改名称")
bullet("启用 / 禁用线条")
body("每个员工和主管账号都属于某条线条，主管只能看本线条成员的数据。")

h2("9.3 审计日志")
add_img(f"{OUT}/07_admin_audit.png", width=Inches(5.8), caption="图 9  审计日志页面")
doc.add_paragraph()
body("记录所有管理员的写操作，包括：建账号、改角色、重置密码、建/改线条、存量认领等。支持按操作人、操作类型、日期范围筛选，不可删除，用于事后审计追责。")

h2("9.4 存量认领")
body("2026 年 5 月 23 日账号模块上线前的历史诊断记录，提交人字段为空（存量数据），仅管理员可见。")
body("管理员可在「存量认领」页面将这些历史记录批量归属到具体账号，方便主管和员工查阅自己线条内的历史数据。")

doc.add_page_break()

# ════════════════════════════════════════════
# 十、常见问题
# ════════════════════════════════════════════
h1("十、常见问题")

h3("Q1：提交诊断后页面长时间没有反应，是否卡住了？")
body("正常现象。系统在规则引擎检测完成后，会对每条触发规则调用 AI 进行个性化分析，通常需要 30～90 秒。请勿刷新页面，等待跳转到报告即可。")
doc.add_paragraph()

h3("Q2：AI 提取的字段有误，如何纠正？")
body("在右侧「信息解析」面板中，点击任意字段可直接修改。修改后继续对话或直接提交均可。提交前系统还会再次确认字段值。")
doc.add_paragraph()

h3("Q3：报告中的「人工核查项目」是什么意思？是说项目一定有问题吗？")
body("不是。人工核查项目是系统无法通过填报字段自动判断的规则（例如「十个不准」），需要用户自行对照每条说明逐项确认。这些项目不计入自动风险等级，是提醒性质的自查清单。")
doc.add_paragraph()

h3("Q4：「硬转服务嫌疑」标记出来就代表违规吗？")
body("不是。硬转服务是举证式检测，系统标记嫌疑并列出需要准备的举证材料，由审核人员根据实际材料综合判断。工具的原则是「标风险、举证定生死，不替审核人定罪」。")
doc.add_paragraph()

h3("Q5：同一项目可以诊断多次吗？")
body("可以。每次提交都是独立的一条诊断记录。可通过「BPM 查询」按商机编号查看同一项目的多次诊断历史，对比风险变化。")
doc.add_paragraph()

h3("Q6：主管能看到下属的诊断记录，下属能看主管的吗？")
body("不能。数据隔离是单向的：主管可以看本线条所有员工的记录，员工只能看自己创建的记录，无法查看主管或同事的数据。")
doc.add_paragraph()

h3("Q7：密码忘了怎么办？")
body("联系管理员，在「账号管理」页面找到你的账号，点击「重置密码」。管理员会给你新的临时密码，你登录后系统会要求立即修改。")
doc.add_paragraph()

h3("Q8：报告下载的 PDF 打不开或下载失败？")
body("确保浏览器没有拦截弹出窗口或下载权限。如问题持续，尝试使用 Chrome 浏览器。PDF 导出功能依赖系统配置，如服务端未安装 WeasyPrint，会自动降级为 HTML 格式下载。")

doc.add_page_break()

# ════════════════════════════════════════════
# 版本信息
# ════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run("─" * 40)
set_font(run, size=10, color=(0xd1, 0xd5, 0xdb))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("规则库版本 v1.7 · 广州电信云中台内部使用 · 2026 年 6 月")
set_font(run, size=9, color=(0x9c, 0xa3, 0xaf))

doc.save(DOC_PATH)
print(f"Word 文档已生成：{DOC_PATH}")
